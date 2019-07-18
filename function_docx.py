#Include Docx files here

import browser_cookie3
import requests
cj = browser_cookie3.chrome()

import re
from docx import Document
def geturl_docx(url):
    url = url.replace('$32', '?')
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}
    fname = requests.get(url, cookies=cj, headers= headers)

    with open('document2.docx', 'wb') as f:
        f.write(fname.content)

    document = Document('document2.docx')
    x = []
    for para in document.paragraphs:
        x.append(para.text)
        for item in x:
            if (len(str(item)) < 1):
                x.remove(str(item))

    document = Document('document2.docx')
    tables = document.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # print(paragraph.text)
                    x.append(paragraph.text)

    urllist = re.compile('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')

    clean_urls = []
    urls = []
    for item in x:
        urls.append( urllist.findall(str(item)))
    for i in urls:
        clean_urls += i
    # print(clean_urls)
    # print(clean_urls)

    import zipfile

    zip_ref = zipfile.ZipFile('document1.docx')
    zip_ref.extractall('document')

    hyper_file = 'document/word/_rels/document.xml.rels'

    file = open(hyper_file,'r')
    xml = file.read()
    file.close()

    xml_urls = urllist.findall(xml)
    # print(xml_urls)
    removelist = set(['http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/package/2006/relationships',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/printerSettings',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing',
        'http://schemas.openxmlformats.org/package/2006/relationships',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme',
        'http://schemas.microsoft.com/office/2007/relationships/stylesWithEffects',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings'
        'http://schemas.openxmlformats.org/package/2006/relationships',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout',
        'http://schemas.openxmlformats.org/package/2006/relationships',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings',
        'http://schemas.microsoft.com/sharepoint/v3/contenttype/forms',
        'http://schemas.microsoft.com/office/2006/metadata/properties/metaAttributes',
        'http://purl.org/dc/elements/1.1/',
        'http://schemas.microsoft.com/office/2006/metadata/longProperties',
        'http://purl.org/dc/terms/',
        'http://www.w3.org/2001/XMLSchema-instance',
        'http://schemas.microsoft.com/office/2006/metadata/properties',
        'http://dublincore.org/schemas/xmls/qdc/2003/04/02/dcterms.xsd',
        'http://schemas.microsoft.com/office/infopath/2007/PartnerControls',
        'http://schemas.microsoft.com/office/2006/documentManagement/types',
        'http://www.w3.org/2001/XMLSchema',
        'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
        'http://schemas.openxmlformats.org/officeDocument/2006/customXml',
        'http://schemas.microsoft.com/internal/obd',
        'http://schemas.openxmlformats.org/drawingml/2006/main',
        'http://dublincore.org/schemas/xmls/qdc/2003/04/02/dc.xsd',
        'http://schemas.microsoft.com/office/2006/metadata/contentType',
        'http://schemas.microsoft.com/office/infopath/2007/Partner<?xml',
        'http://schemas.microsoft.com/sharepoint/events',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/package',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
        'http://www.w3.org/1999/02/22-rdf-syntax-ns',
        'http://schemas.openxmlformats.org/officeDocument/2006/bibliography','http://schemas.openxmlformats.org/package/2006/relationships',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer', 'http://schemas.microsoft.com/office/2007/relationships/stylesWithEffects', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',"http://<?xml",
        "http://schemas.microsoft.com/office/<?xml",])
    true = []
    for item in xml_urls:
        if item not in removelist:
            true.append(item)

    # print(true)
    for u in true:
        clean_urls.append(u)
    # print(clean_urls)
    return clean_urls
    # print(urls)
    # for item in urls:
    #     if (len(item) < 1):
    #         urls.remove(item)
    # print(urls)
    # total = []
    # for i in urls:
    #     total += i
    # print(total)


    # true = []
    # # for item in urls:
    # #     true += item
    # print(true)

    # print(clean_urls)
    # for j in urls:
    #     if j not in removelist:
    #         true.append(j)
    # return true
    # print(true)

